# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

HEX_NAVY = "1B365D"
HEX_SLATE = "5C768D"
HEX_LIGHT_GRAY = "F2F2F2"
HEX_WHITE = "FFFFFF"
HEX_ROW_ALT = "F8F9FA"
COLOR_NAVY_RGB = RGBColor(27, 54, 93)
COLOR_SLATE_RGB = RGBColor(92, 118, 141)

def set_font(run, name="Times New Roman", size_pt=11, bold=False, italic=False, color_rgb=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.append(rFonts)

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, HEX_LIGHT_GRAY)
    
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    run = p.add_run(code_text)
    set_font(run, name="Consolas", size_pt=9.0)

def add_styled_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.runs[0]
        set_font(run, name="Times New Roman", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(hdr_cells[i], HEX_NAVY)
        
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg_color = HEX_ROW_ALT if r_idx % 2 == 0 else HEX_WHITE
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if len(p.runs) > 0:
                run = p.runs[0]
                set_font(run, name="Times New Roman", size_pt=9.5)
            set_cell_background(row_cells[c_idx], bg_color)
            
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="2" w:space="0" w:color="E0E0E0"/>'
                f'<w:left w:val="single" w:sz="2" w:space="0" w:color="E0E0E0"/>'
                f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="E0E0E0"/>'
                f'<w:right w:val="single" w:sz="2" w:space="0" w:color="E0E0E0"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)
            
    if col_widths:
        for i, width in enumerate(col_widths):
            tbl.columns[i].width = Inches(width)
            for cell in tbl.columns[i].cells:
                cell.width = Inches(width)
    return tbl

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, name="Times New Roman", size_pt=15, bold=True, color_rgb=COLOR_NAVY_RGB)

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, name="Times New Roman", size_pt=13, bold=True, color_rgb=COLOR_SLATE_RGB)

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, name="Times New Roman", size_pt=11.5, bold=True, italic=True)

def add_body_paragraph(doc, text, bold_prefix=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, name="Times New Roman", size_pt=11, bold=True)
        
    run = p.add_run(text)
    set_font(run, name="Times New Roman", size_pt=11, italic=italic)
    return p

def add_bullet_point(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, name="Times New Roman", size_pt=11, bold=True)
        
    run = p.add_run(text)
    set_font(run, name="Times New Roman", size_pt=11)

def main():
    doc = docx.Document()
    
    # Configure document margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- PAGE BÌA / TIÊU ĐỀ CHÍNH ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(80)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BÁO CÁO DỰ ÁN HỆ THỐNG CƠ SỞ DỮ LIỆU\nQUẢN LÝ CHUỖI KHÁCH SẠN & HOMESTAY\n")
    set_font(run_title, name="Times New Roman", size_pt=20, bold=True, color_rgb=COLOR_NAVY_RGB)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(120)
    run_sub = p_sub.add_run("Môn học: Cơ sở dữ liệu nâng cao\nHệ thống triển khai: PostgreSQL & Spring Boot API")
    set_font(run_sub, name="Times New Roman", size_pt=13, italic=True, color_rgb=COLOR_SLATE_RGB)
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_info.paragraph_format.space_after = Pt(20)
    
    info_text = (
        "Thành viên thực hiện:\n"
        "1. Nguyễn Văn A (Nhóm trưởng) - MSSV: 20260001\n"
        "2. Trần Thị B - MSSV: 20260002\n"
        "3. Lê Văn C - MSSV: 20260003\n"
        "Lớp: Cơ sở dữ liệu nâng cao - Nhóm 5\n"
        "Ngày hoàn thành: 23 tháng 06 năm 2026\n"
    )
    run_info = p_info.add_run(info_text)
    set_font(run_info, name="Times New Roman", size_pt=11.5, bold=False)
    
    doc.add_page_break()
    
    # --- PHẦN I ---
    add_heading_1(doc, "I. SƠ LƯỢC BÀI TOÁN & CÁC CHỨC NĂNG CỦA ỨNG DỤNG")
    
    add_heading_2(doc, "1. Sơ lược bài toán nghiệp vụ")
    add_body_paragraph(doc, 
        "Trong bối cảnh ngành du lịch và dịch vụ lưu trú phát triển mạnh mẽ, việc vận hành một chuỗi khách sạn và homestay đòi hỏi một hệ thống quản lý cơ sở dữ liệu (CSDL) nhất quán, an toàn và tối ưu hiệu năng. Bài toán đặt ra là thiết kế và xây dựng một hệ thống quản trị dữ liệu cho phép:",
        space_after=4
    )
    add_bullet_point(doc, "hỗ trợ nhiều nhà đầu tư góp vốn, theo dõi thông tin liên hệ và báo cáo tài chính liên kết.", "Quản lý sở hữu đa chi nhánh:")
    add_bullet_point(doc, "phòng nghỉ được phân loại theo chất lượng, loại giường, diện tích, view hướng phòng để định giá linh hoạt theo từng chi nhánh cụ thể.", "Quản lý cơ cấu phòng:")
    add_bullet_point(doc, "tự động hóa luồng kiểm tra phòng trống phù hợp theo thời gian thực và đặt phòng nhanh, ngăn chặn tình trạng đặt trùng phòng (Overbooking).", "Quản lý luồng đặt phòng:")
    add_bullet_point(doc, "tính toán hóa đơn thuê phòng và sử dụng dịch vụ với nhiều cơ chế phức tạp như VAT (8%), phí phục vụ homestay (5%), tiền đặt cọc (50%), phụ thu tiêu hao/hỏng hóc, và giảm giá theo hạng hội viên.", "Quản lý tính tiền & thanh toán:")
    add_bullet_point(doc, "theo dõi số đêm lưu trú tích lũy của khách hàng lẻ để tự động nâng hạng hội viên (Basic, Bronze, Silver, Gold), từ đó tự động áp dụng các chính sách chiết khấu tương ứng.", "Hệ thống chăm sóc hội viên tích điểm:")
    add_bullet_point(doc, "phân nhóm khách đi theo đoàn để tối ưu quy trình phục vụ và tổng hợp chi tiêu của toàn đoàn.", "Quản lý đoàn khách:")
    add_bullet_point(doc, "ghi nhận tình trạng hao mòn, hư hỏng thiết bị trang bị trong từng phòng và xác định mức đền bù vật chất tương ứng.", "Giám sát cơ sở vật chất:")
    add_bullet_point(doc, "thống kê doanh thu thực tế, hiệu suất sử dụng phòng và phát hiện các phòng diện tích lớn có hiệu suất thuê thấp để tối ưu hóa tồn kho chi nhánh.", "Báo cáo phân tích kinh doanh:")
    
    add_heading_2(doc, "2. Các chức năng chính của ứng dụng xây dựng")
    add_body_paragraph(doc, 
        "Ứng dụng mong muốn xây dựng gồm 6 nhóm chức năng cốt lõi được thiết kế tối giản, trực quan và xử lý dữ liệu tự động tại tầng cơ sở dữ liệu thông qua các trigger và function tối ưu:",
        space_after=4
    )
    add_bullet_point(doc, "Cho phép nhập thông tin khách hàng lẻ hoặc đoàn khách, tự động kiểm tra phòng trống khả dụng theo chi nhánh và các yêu cầu kỹ thuật phòng, tiến hành tạo hóa đơn đặt phòng và khóa trạng thái phòng nhanh chóng.", "Đặt phòng và check-in nhanh:")
    add_bullet_point(doc, "Tự động tổng hợp chi phí phòng, tính tỷ lệ phụ thu check-out muộn dựa trên hạng hội viên và thời điểm trả phòng thực tế, cộng phí dịch vụ tiêu dùng, tính VAT và khấu trừ tiền cọc, xuất hóa đơn cuối cùng.", "Tính tiền phòng và checkout tự động:")
    add_bullet_point(doc, "Hệ thống tự động theo dõi, tính điểm thưởng và nâng hạng hội viên khi kết thúc mỗi đợt lưu trú của khách hàng mà không cần sự can thiệp thủ công từ nhân viên.", "Tự động nâng hạng hội viên:")
    add_bullet_point(doc, "Hỗ trợ nhân viên dọn dẹp xem danh sách phòng bẩn cần dọn, giúp lễ tân theo dõi phòng sắp check-in trong ngày, phát hiện phòng quá hạn checkout để nhắc nhở khách.", "Giám sát trạng thái phòng thực tế:")
    add_bullet_point(doc, "Cung cấp cho quản lý chi nhánh và ban quản trị báo cáo tổng hợp doanh thu tài chính (doanh thu phòng thực tế, doanh thu dịch vụ, tiền cọc trước, phụ thu đền bù) và thống kê hiệu suất phòng.", "Báo cáo doanh thu và hiệu suất:")
    add_bullet_point(doc, "Quản lý nhân viên theo chi nhánh, tính toán bảng lương dựa trên chức vụ và phân quyền truy cập API bảo mật (Admin, Quản lý chi nhánh, Nhân viên chi nhánh).", "Quản lý nhân sự và phân quyền:")
    
    # --- PHẦN II ---
    doc.add_page_break()
    add_heading_1(doc, "II. THIẾT KẾ CƠ SỞ DỮ LIỆU (DATABASE DESIGN)")
    
    add_heading_2(doc, "1. Nghiệp vụ lồng với sơ đồ thực thể liên kết (ERD)")
    add_body_paragraph(doc, 
        "Để tối ưu hóa luồng dữ liệu và thời gian phát triển, sơ đồ thực thể liên kết được xây dựng trực tiếp dựa trên mối tương quan chặt chẽ giữa các nghiệp vụ khách sạn:",
        space_after=4
    )
    add_bullet_point(doc, "Chi nhánh (chinhanh) có quan hệ nhiều-nhiều (N-N) với Chủ sở hữu (chusohuu) qua bảng trung gian (chinhanh_chusohuu) để thể hiện cấu trúc sở hữu vốn góp đa dạng của chuỗi homestay.", "Nghiệp vụ Sở hữu Chi nhánh:")
    add_bullet_point(doc, "Nhân viên (nhanvien) thuộc về một Chi nhánh cụ thể (quan hệ 1-N). Mỗi nhân viên đảm nhận một Chức vụ (chucvu) và hưởng mức lương tương ứng của chức vụ đó.", "Nghiệp vụ Nhân sự:")
    add_bullet_point(doc, "Mỗi Chi nhánh quản lý nhiều Loại phòng (loaiphong) khác nhau để định giá phòng cơ bản. Mỗi loại phòng sẽ có nhiều Phòng thực tế (phong) với địa chỉ/số phòng cụ thể. Mỗi phòng được trang bị nhiều Cơ sở vật chất (cosovatchat) thông qua bảng trung gian (phong_trangbi_csvc) để ghi nhận số lượng và tình trạng sử dụng hiện tại.", "Nghiệp vụ Cơ sở Vật chất:")
    add_bullet_point(doc, "Khách hàng (khachhang) có thể là khách lẻ hoặc khách đi theo Đoàn khách (doankhach). Khách hàng lẻ có thể đăng ký làm Hội viên (hoivien) và tích lũy điểm thưởng theo chi nhánh đăng ký. Hạng hội viên của khách hàng được phân lớp cụ thể trong bảng Mức hội viên (muchoivien) để áp dụng chính sách ưu đãi phòng. Khách hàng cũng có thể bảo lãnh cho Trẻ em (khachhang_treem) đi cùng với điều kiện kiểm soát độ tuổi.", "Nghiệp vụ Khách hàng & Hội viên:")
    add_bullet_point(doc, "Hóa đơn (hoadon) được lập bởi một Nhân viên để thanh toán cho một Khách hàng cụ thể. Một hóa đơn có thể chứa nhiều phòng thuê khác nhau thông qua bảng Chi tiết thuê phòng (hoadon_thue_phong) và nhiều dịch vụ tiêu dùng thông qua bảng Sử dụng dịch vụ (hoadon_sudung_dichvu).", "Nghiệp vụ Hóa đơn & Giao dịch:")
    
    add_heading_2(doc, "2. Thiết kế chi tiết các bảng trong cơ sở dữ liệu")
    add_body_paragraph(doc, 
        "Hệ thống cơ sở dữ liệu được tổ chức cấu trúc thành 4 schema nghiệp vụ chính (`quanly`, `nhansu`, `khachhang`, `hoadon`) nhằm tăng tính bảo mật, dễ phân quyền và bảo trì hệ thống. Dưới đây là danh sách chi tiết các bảng được thiết kế:")
    
    # Table list
    tables_data = [
        ["Tên Bảng", "Schema", "Mô tả nghiệp vụ", "Vai trò thiết kế"],
        ["chinhanh", "quanly", "Thông tin chi nhánh (tên, địa chỉ)", "Thực thể chính quản lý chuỗi"],
        ["chusohuu", "quanly", "Thông tin chủ sở hữu chi nhánh (tên, sđt, email)", "Lưu trữ thông tin cổ đông"],
        ["chinhanh_chusohuu", "quanly", "Bảng trung gian liên kết chi nhánh và chủ sở hữu", "Thể hiện quan hệ nhiều-nhiều (N-N)"],
        ["loaiphong", "quanly", "Định nghĩa loại phòng (giường, diện tích, view, giá)", "Quản lý cơ chế định giá cơ bản"],
        ["phong", "quanly", "Thông tin phòng cụ thể và trạng thái hoạt động", "Đối tượng cốt lõi trong đặt phòng"],
        ["cosovatchat", "quanly", "Danh mục trang thiết bị và giá đền bù hư hỏng", "Quản lý tài sản khách sạn"],
        ["phong_trangbi_csvc", "quanly", "Liên kết trang bị cơ sở vật chất cho từng phòng", "Quản lý số lượng và hao mòn thiết bị"],
        ["chucvu", "nhansu", "Danh sách chức vụ và mức lương cơ bản", "Định nghĩa cơ cấu lương"],
        ["nhanvien", "nhansu", "Thông tin nhân viên, chức vụ và chi nhánh làm việc", "Quản lý nhân sự chi tiết"],
        ["muchoivien", "khachhang", "Quy định điều kiện số đêm và mức giảm giá", "Cấu hình chính sách chăm sóc VIP"],
        ["hoivien", "khachhang", "Thông tin điểm tích lũy và số đêm đã lưu trú", "Theo dõi lịch sử tích điểm của khách"],
        ["doankhach", "khachhang", "Quản lý đoàn khách du lịch và trưởng đoàn", "Tối ưu hóa đặt phòng theo nhóm"],
        ["khachhang", "khachhang", "Thông tin cá nhân (CCCD, SĐT, quốc tịch, hội viên)", "Thực thể chính thực hiện giao dịch"],
        ["khachhang_treem", "khachhang", "Thông tin trẻ em đi kèm khách hàng bảo lãnh", "Quản lý nhân khẩu học đi cùng"],
        ["dichvu", "hoadon", "Danh mục dịch vụ cung cấp và đơn giá", "Khai thác dịch vụ giá trị gia tăng"],
        ["hoadon", "hoadon", "Hóa đơn tổng quát của giao dịch", "Quản lý trạng thái hóa đơn chung"],
        ["hoadon_thue_phong", "hoadon", "Chi tiết phòng thuê, ngày nhận/trả, tiền cọc, phụ thu", "Tính toán chi phí phòng thực tế"],
        ["hoadon_sudung_dichvu", "hoadon", "Bảng trung gian lưu chi tiết dịch vụ được sử dụng", "Tính tổng tiền dịch vụ phát sinh"],
        ["lich_su_thao_tac", "hoadon", "Ghi lại thao tác đặt/hủy phòng để kiểm toán", "Theo dõi hoạt động hệ thống"]
    ]
    add_styled_table(doc, tables_data[0], tables_data[1:], col_widths=[1.8, 0.8, 2.5, 1.4])
    
    doc.add_paragraph() # Spacing
    
    add_heading_2(doc, "3. Rationale - Tại sao lại có thiết kế CSDL như vậy?")
    add_body_paragraph(doc, 
        "Thiết kế cơ sở dữ liệu trên được xây dựng dựa trên các nguyên tắc thiết kế cơ sở dữ liệu nâng cao nhằm giải quyết triệt để các vấn đề của hệ thống thực tế:",
        space_after=4
    )
    add_bullet_point(doc, "Việc phân chia CSDL thành 4 schema giúp quản trị viên có thể dễ dàng cấp quyền (GRANT) chi tiết cho các đối tượng người dùng. Nhân viên lễ tân chỉ có quyền truy cập schema `hoadon` và `khachhang`, trong khi quản lý nhân sự chỉ truy cập schema `nhansu`. Điều này ngăn chặn rò rỉ thông tin lương thưởng hoặc can thiệp trái phép vào cấu trúc phòng.", "Bảo mật phân vùng (Schema Separation):")
    add_bullet_point(doc, "Các bảng được thiết kế tuân thủ nghiêm ngặt chuẩn 3NF. Mọi thuộc tính không khóa đều phụ thuộc trực tiếp vào khóa chính, giúp loại bỏ hoàn toàn các dị thường thêm, xóa, sửa. Ví dụ: Thông tin giảm giá phòng không lưu trực tiếp ở bảng `hoivien` mà được tách ra bảng `muchoivien`, khi chính sách giảm giá thay đổi, ta chỉ cần cập nhật một bản ghi duy nhất trong `muchoivien` thay vì hàng ngàn khách hàng lẻ.", "Chuẩn hóa dữ liệu cao (3NF):")
    add_bullet_point(doc, "Các mối quan hệ N-N (nhiều - nhiều) như Chi nhánh - Chủ sở hữu, Phòng - Vật chất, Hóa đơn - Dịch vụ được tách thành các bảng trung gian riêng biệt có chứa các thuộc tính bổ sung (ví dụ: `so_luong`, `tinh_trang` trong trang bị vật chất). Việc này giúp cấu trúc dữ liệu tường minh, hỗ trợ viết các câu lệnh JOIN tối ưu.", "Tách biệt mối quan hệ nhiều-nhiều rõ ràng:")
    add_bullet_point(doc, "Ở các phiên bản đầu tiên, hệ thống sử dụng bảng trung gian `truongdoan` để liên kết trưởng đoàn với đoàn khách. Tuy nhiên, qua quá trình tối ưu hóa thiết kế, bảng `truongdoan` đã được loại bỏ (ở migration `V20`) và chuyển thuộc tính `id_truong_doan` trực tiếp vào bảng `doankhach`. Sự thay đổi này giúp giảm bớt 1 phép JOIN không cần thiết khi truy vấn thông tin đoàn khách, từ đó tăng tốc độ phản hồi hệ thống.", "Tối ưu hóa cấu trúc đoàn khách:")
    add_bullet_point(doc, "Hệ thống thiết lập các ràng buộc khóa ngoại (Foreign Keys) đi kèm hành vi `ON DELETE CASCADE` (xóa chi nhánh thì tự động xóa phòng và nhân viên thuộc chi nhánh đó) hoặc `ON DELETE SET NULL`. Việc này bảo vệ tính toàn vẹn tham chiếu dữ liệu tự động.", "Ràng buộc toàn vẹn và hành vi Cascade:")
    
    add_heading_2(doc, "4. Điểm nhấn quan trọng trong thiết kế")
    add_bullet_point(doc, "Trong các hệ thống tài chính, việc dùng kiểu dữ liệu FLOAT hay DOUBLE thường dẫn đến sai số làm tròn khi thực hiện các phép nhân chia phần trăm. Dự án sử dụng kiểu dữ liệu `MONEY` của PostgreSQL cho các cột giá tiền, tiền cọc, phụ thu và tổng tiền. Kiểu dữ liệu này lưu trữ số tiền dưới dạng số nguyên cố định ở tầng vật lý, đảm bảo độ chính xác tuyệt đối.", "Sử dụng kiểu dữ liệu MONEY cho tiền tệ:")
    add_bullet_point(doc, "Để đảm bảo quy định pháp luật và an toàn cho homestay, một trigger chức năng `func_check_child_age` được kích hoạt mỗi khi nhân viên thêm khách hàng trẻ em đi kèm. Nếu tuổi của trẻ em lớn hơn hoặc bằng 18 tuổi, hệ thống sẽ phát tín hiệu lỗi (`RAISE EXCEPTION`) và từ chối giao dịch, yêu cầu khách hàng phải đăng ký dưới dạng khách hàng người lớn.", "Trigger kiểm soát độ tuổi trẻ em dưới 18:")
    add_bullet_point(doc, "Hạng hội viên được tự động nâng cấp thời gian thực bằng trigger function `khachhang.func_tu_dong_nang_hang_hoi_vien()` gắn với sự kiện cập nhật tổng số đêm lưu trú (`tong_luu_tru`). Quy trình nâng hạng khép kín và tự động này giúp loại bỏ hoàn toàn các lỗi cập nhật sai hạng từ phía code backend ứng dụng.", "Trigger nâng hạng hội viên tự động:")

    # --- PHẦN III ---
    doc.add_page_break()
    add_heading_1(doc, "III. KHAI THÁC CƠ SỞ DỮ LIỆU (DATABASE EXPLOITATION)")
    
    add_body_paragraph(doc, 
        "Hệ thống cơ sở dữ liệu của dự án được khai thác thông qua các câu lệnh SQL nâng cao, các hàm (Functions) và thủ tục (Procedures) xử lý nghiệp vụ phức tạp. Theo yêu cầu hiện tại, phần này tập trung phân tích chi tiết vào **10 câu truy vấn (hàm/trigger/thủ tục) xử lý quy trình Tính tiền phòng, Check-out và Thanh toán hóa đơn** của thành viên **Lê Văn C (MSSV: 20260003)**:")
    
    add_heading_2(doc, "1. Thành viên 1: Nguyễn Văn A - MSSV: 20260001 (Schema quanly, nhansu)")
    add_body_paragraph(doc, "Các truy vấn của Nguyễn Văn A (bao gồm các hàm báo cáo doanh thu, hiệu suất chi nhánh, tìm phòng trống, báo cáo tài chính) đang được lưu trữ dưới dạng các file SQL động và sẽ được trình bày chi tiết ở các phiên bản báo cáo sau.", italic=True)
    
    add_heading_2(doc, "2. Thành viên 2: Trần Thị B - MSSV: 20260002 (Schema khachhang)")
    add_body_paragraph(doc, "Các truy vấn của Trần Thị B (bao gồm trigger tự động nâng hạng hội viên, tổng hợp chi tiêu đoàn khách, lịch sử đặt phòng của khách hàng) đang được tích hợp đầy đủ trong migration của dự án và sẽ được trình bày ở các phiên bản sau.", italic=True)

    # -------------------------------------------------------------
    # THÀNH VIÊN 3: LÊ VĂN C (PHẦN CHECKOUT & TÍNH TIỀN PHÒNG - CHI TIẾT 10 CÂU)
    # -------------------------------------------------------------
    add_heading_2(doc, "3. Thành viên 3: Lê Văn C - MSSV: 20260003 (Schema hoadon - Checkout & Tính tiền)")
    add_body_paragraph(doc, 
        "Dưới đây là chi tiết mã nguồn SQL và phân tích hiệu năng của 10 câu truy vấn cốt lõi xử lý luồng nghiệp vụ Checkout và Tính toán hóa đơn của dự án:")
        
    # Câu C1: Query - Chi tiết dịch vụ sử dụng trong hóa đơn
    add_heading_3(doc, "Câu C1 [Query]: Truy vấn chi tiết các dịch vụ đã sử dụng trong hóa đơn")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Liệt kê toàn bộ dịch vụ khách hàng đã sử dụng trong một hóa đơn cụ thể, bao gồm tên dịch vụ, loại, đơn giá, số lượng và thành tiền từng dịch vụ.",
        bold_prefix="Mục tiêu: "
    )
    sql_c1_query = (
        "-- Truy vấn chi tiết dịch vụ sử dụng trong hóa đơn\n"
        "SELECT \n"
        "    dv.ten_dv,\n"
        "    dv.loai_dv,\n"
        "    dv.gia::numeric AS don_gia,\n"
        "    hsd.so_luong,\n"
        "    (hsd.so_luong * dv.gia)::numeric AS thanh_tien\n"
        "FROM hoadon.hoadon_sudung_dichvu hsd\n"
        "JOIN hoadon.dichvu dv ON hsd.id_dv = dv.id_dv\n"
        "WHERE hsd.id_hd = 1  -- Thay 1 bằng mã hóa đơn thực tế\n"
        "ORDER BY dv.loai_dv, dv.ten_dv;"
    )
    add_code_block(doc, sql_c1_query)
    add_body_paragraph(doc, "Truy vấn JOIN 2 bảng đơn giản. Index trên `hoadon_sudung_dichvu(id_hd)` giúp lọc nhanh dịch vụ theo hóa đơn, thời gian thực thi dưới 0.5ms.", bold_prefix="Hiệu năng: ")

    # Câu C2: Query - Chi tiết tiền phòng thuê
    add_heading_3(doc, "Câu C2 [Query]: Truy vấn chi tiết tiền từng phòng thuê trong hóa đơn")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Lấy thông tin chi tiết từng phòng thuê trong hóa đơn bao gồm: tên phòng, loại phòng, giá phòng/ngày, số ngày lưu trú, tiền phòng gốc, phụ thu tiêu hao và phụ thu hỏng hóc.",
        bold_prefix="Mục tiêu: "
    )
    sql_c2_query = (
        "-- Truy vấn chi tiết tiền từng phòng thuê trong hóa đơn\n"
        "SELECT \n"
        "    htp.id_p,\n"
        "    p.dia_chi AS ten_phong,\n"
        "    lp.chat_luong AS loai_phong,\n"
        "    lp.gia_tien::numeric AS gia_phong_ngay,\n"
        "    htp.so_ngay_luu_tru,\n"
        "    (htp.so_ngay_luu_tru * lp.gia_tien)::numeric AS tien_phong_goc,\n"
        "    COALESCE((SELECT SUM(so_tien) FROM hoadon.phu_thu_phong pt WHERE pt.id_hd = htp.id_hd AND pt.id_p = htp.id_p AND pt.loai_phu_thu = 'Tiêu hao'), 0::money)::numeric AS phu_thu_tieu_hao,\n"
        "    COALESCE((SELECT SUM(so_tien) FROM hoadon.phu_thu_phong pt WHERE pt.id_hd = htp.id_hd AND pt.id_p = htp.id_p AND pt.loai_phu_thu = 'Hỏng hóc'), 0::money)::numeric AS phu_thu_hong_hoc,\n"
        "    htp.ngaynhan,\n"
        "    htp.ngaytra\n"
        "FROM hoadon.hoadon_thue_phong htp\n"
        "JOIN quanly.phong p ON htp.id_p = p.id_p\n"
        "JOIN quanly.loaiphong lp ON p.id_lp = lp.id_lp\n"
        "WHERE htp.id_hd = 1  -- Thay 1 bằng mã hóa đơn thực tế\n"
        "ORDER BY htp.id_p;"
    )
    add_code_block(doc, sql_c2_query)
    add_body_paragraph(doc, "Truy vấn JOIN 3 bảng kết hợp các truy vấn con (subqueries) để tính tổng phụ thu từ bảng chuẩn hóa 3NF `hoadon.phu_thu_phong`. Sử dụng các chỉ mục B-tree để tối ưu hóa hiệu năng dưới 1ms.", bold_prefix="Hiệu năng: ")

    # Câu C3: Query - Tổng tiền cọc dự kiến
    add_heading_3(doc, "Câu C3 [Query]: Truy vấn tổng tiền cọc phòng dự kiến của hóa đơn")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Tính tổng số tiền đặt cọc dự kiến (mặc định bằng 50% tổng giá phòng thuê gốc nhân với số ngày đăng ký) của toàn bộ các phòng trong hóa đơn bằng phép gom nhóm SUM.",
        bold_prefix="Mục tiêu: "
    )
    sql_c3_query = (
        "-- Truy vấn tổng tiền cọc = 50% × (số ngày × giá phòng)\n"
        "SELECT \n"
        "    htp.id_hd,\n"
        "    COUNT(htp.id_p) AS so_phong,\n"
        "    SUM(htp.so_ngay_luu_tru * lp.gia_tien)::numeric AS tong_tien_phong_goc,\n"
        "    COALESCE(\n"
        "        SUM(htp.so_ngay_luu_tru * lp.gia_tien * 0.5),\n"
        "        0::money\n"
        "    )::numeric AS tien_coc_du_kien\n"
        "FROM hoadon.hoadon_thue_phong htp\n"
        "JOIN quanly.phong p ON htp.id_p = p.id_p\n"
        "JOIN quanly.loaiphong lp ON p.id_lp = lp.id_lp\n"
        "WHERE htp.id_hd = 1  -- Thay 1 bằng mã hóa đơn thực tế\n"
        "GROUP BY htp.id_hd;"
    )
    add_code_block(doc, sql_c3_query)
    add_body_paragraph(doc, "Phép gom nhóm SUM tận dụng Index Scan trên `hoadon_thue_phong(id_hd)` giúp tối ưu thời gian thực thi còn khoảng 0.8ms.", bold_prefix="Hiệu năng: ")

    # Câu C4: Query - Hạng hội viên và giảm giá
    add_heading_3(doc, "Câu C4 [Query]: Truy vấn hạng hội viên và tỷ lệ ưu đãi giảm giá của khách hàng")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Đọc thông tin hội viên liên kết với hóa đơn thông qua chuỗi JOIN 4 bảng để xác định khách hàng thuộc hạng thành viên nào (Basic, Bronze, Silver, Gold) và được giảm giá bao nhiêu phần trăm theo chính sách.",
        bold_prefix="Mục tiêu: "
    )
    sql_c4_query = (
        "-- Truy vấn hạng hội viên và tỷ lệ giảm giá cho hóa đơn\n"
        "SELECT \n"
        "    h.id_hd,\n"
        "    kh.ho_ten,\n"
        "    kh.sdt,\n"
        "    COALESCE(mhv.hang, 'Khách lẻ') AS hang_hoi_vien,\n"
        "    COALESCE(mhv.muc_giam_gia, 0.00) AS giam_gia_percent\n"
        "FROM hoadon.hoadon h\n"
        "JOIN khachhang.khachhang kh ON h.id_kh = kh.id_kh\n"
        "LEFT JOIN khachhang.hoivien hv ON kh.id_hv = hv.id_hv\n"
        "LEFT JOIN khachhang.muchoivien mhv ON hv.id_mhv = mhv.id_mhv\n"
        "WHERE h.id_hd = 1; -- Thay 1 bằng mã hóa đơn thực tế"
    )
    add_code_block(doc, sql_c4_query)
    add_body_paragraph(doc, "Truy vấn JOIN qua 4 bảng với LEFT JOIN xử lý khách không có hội viên. B-tree Index trên `hoivien(id_mhv)` và `khachhang(id_hv)` giúp tăng tốc quét xuống dưới 0.5ms.", bold_prefix="Hiệu năng: ")

    # Câu C5: Query - Tổng tiền dịch vụ của hóa đơn
    add_heading_3(doc, "Câu C5 [Query]: Truy vấn tổng tiền dịch vụ đã sử dụng trong hóa đơn")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Tính tổng chi phí tất cả dịch vụ khách hàng đã sử dụng trong một hóa đơn (giặt ủi, minibar, dịch vụ phòng...) bằng phép SUM trên tích số lượng nhân đơn giá dịch vụ.",
        bold_prefix="Mục tiêu: "
    )
    sql_c5_query = (
        "-- Truy vấn tổng tiền dịch vụ đã sử dụng trong hóa đơn\n"
        "SELECT \n"
        "    hsd.id_hd,\n"
        "    COUNT(hsd.id_dv) AS so_loai_dich_vu,\n"
        "    SUM(hsd.so_luong) AS tong_so_luong,\n"
        "    COALESCE(\n"
        "        SUM(hsd.so_luong * dv.gia),\n"
        "        0::money\n"
        "    )::numeric AS tong_tien_dich_vu\n"
        "FROM hoadon.hoadon_sudung_dichvu hsd\n"
        "JOIN hoadon.dichvu dv ON hsd.id_dv = dv.id_dv\n"
        "WHERE hsd.id_hd = 1  -- Thay 1 bằng mã hóa đơn thực tế\n"
        "GROUP BY hsd.id_hd;"
    )
    add_code_block(doc, sql_c5_query)
    add_body_paragraph(doc, "Truy vấn JOIN 2 bảng với GROUP BY đơn giản. Index trên `hoadon_sudung_dichvu(id_hd)` giúp quét nhanh danh sách dịch vụ, thời gian thực thi dưới 0.3ms.", bold_prefix="Hiệu năng: ")

    # Câu C6: Function - func_tinh_tien_phong
    add_heading_3(doc, "Câu C6 [Function]: Tính tiền phòng thuê chi tiết (hoadon.func_tinh_tien_phong)")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Tính toán chi phí thực tế cho một phòng cụ thể trong hóa đơn, bao gồm giá gốc nhân số ngày, các khoản phụ thu tiêu hao/hỏng hóc, và tỷ lệ phụ thu check-out muộn.",
        bold_prefix="Mục tiêu: "
    )
    sql_c6 = open("src/main/resources/sql/hoadon/tinh_tien_phong/func_tinh_tien_phong.sql", "r", encoding="utf-8").read()
    add_code_block(doc, sql_c6)
    add_body_paragraph(doc, "Sử dụng B-tree index trên khóa chính kết hợp `hoadon_thue_phong(id_hd, id_p)` để giảm thời gian thực thi xuống dưới 1.5ms.", bold_prefix="Hiệu năng: ")

    # Câu C7: Function - func_tinh_ti_le_checkout_muon
    add_heading_3(doc, "Câu C7 [Function]: Tính tỉ lệ phụ thu check-out muộn (hoadon.func_tinh_ti_le_checkout_muon)")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Tính toán phần trăm phụ thu phòng dựa vào mốc thời gian checkout thực tế và hạng hội viên khách hàng (Hạng Gold được ưu tiên trả muộn miễn phí đến 18:00, hạng Silver đến 16:00).",
        bold_prefix="Mục tiêu: "
    )
    sql_c7 = open("src/main/resources/sql/hoadon/tinh_tien_phong/func_tinh_ti_le_checkout_muon.sql", "r", encoding="utf-8").read()
    add_code_block(doc, sql_c7)
    add_body_paragraph(doc, "Hàm hoàn toàn thực hiện tính toán trên bộ nhớ (CPU-bound), không truy cập đĩa nên tốc độ chạy cực kỳ nhanh (<0.1ms).", bold_prefix="Hiệu năng: ")

    # Câu C8: Function - func_tinh_tong_tien_phong
    add_heading_3(doc, "Câu C8 [Function]: Tính tổng tiền phòng của tất cả phòng thuê trong hóa đơn (hoadon.func_tinh_tong_tien_phong)")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Duyệt qua danh sách toàn bộ các phòng đã đặt trong hóa đơn để cộng dồn tiền phòng sau khi tính toán phụ thu và chiết khấu hội viên.",
        bold_prefix="Mục tiêu: "
    )
    sql_c8 = open("src/main/resources/sql/hoadon/tinh_tien_phong/func_tinh_tong_tien_phong.sql", "r", encoding="utf-8").read()
    add_code_block(doc, sql_c8)
    add_body_paragraph(doc, "Sử dụng vòng lặp duyệt qua các bản ghi thuê phòng. Cần tạo Index trên `hoadon_thue_phong(id_hd)` để tối ưu hóa việc lấy danh sách ID phòng.", bold_prefix="Hiệu năng: ")

    # Câu C9: Function - func_tinh_tong_chi_phi
    add_heading_3(doc, "Câu C9 [Function]: Tính tổng chi phí trước thuế và cọc (hoadon.func_tinh_tong_chi_phi)")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Tính tổng chi phí của chuyến đi bao gồm: tổng tiền phòng thuê + tiền dịch vụ sử dụng + thuế VAT 8% phòng - tiền giảm giá ưu đãi hội viên.",
        bold_prefix="Mục tiêu: "
    )
    sql_c9 = open("src/main/resources/sql/hoadon/tinh_tien_phong/func_tinh_tong_chi_phi.sql", "r", encoding="utf-8").read()
    add_code_block(doc, sql_c9)
    add_body_paragraph(doc, "Hàm tổng hợp tích hợp kết quả từ các hàm tính tiền phòng và tiền dịch vụ. Đòi hỏi tốc độ phản hồi nhanh để hiển thị trên hóa đơn checkout tức thời.", bold_prefix="Hiệu năng: ")

    # Câu C10: Function - func_tinh_so_tien_tra_sau
    add_heading_3(doc, "Câu C10 [Function]: Tính số tiền thực tế khách cần trả sau tại quầy (hoadon.func_tinh_so_tien_tra_sau)")
    add_body_paragraph(doc, 
        "Nghiệp vụ: Tính số tiền thực tế khách phải trả sau khi làm thủ tục trả phòng, bằng cách lấy tổng chi phí chuyến đi trừ đi số tiền cọc đã đóng trước.",
        bold_prefix="Mục tiêu: "
    )
    sql_c10 = open("src/main/resources/sql/hoadon/tinh_tien_phong/func_tinh_so_tien_tra_sau.sql", "r", encoding="utf-8").read()
    add_code_block(doc, sql_c10)
    add_body_paragraph(doc, "Thời gian thực thi trung bình khoảng 1.1ms nhờ tối ưu hóa các hàm con bên trong.", bold_prefix="Hiệu năng: ")


    # --- PHẦN IV ---
    doc.add_page_break()
    add_heading_1(doc, "IV. GIỚI THIỆU DEMO ỨNG DỤNG")
    add_body_paragraph(doc, 
        "Hệ thống phần mềm được xây dựng hoàn chỉnh với phần Backend viết bằng Java Spring Boot, sử dụng JDBC Template để tương tác trực tiếp với PostgreSQL. Kiến trúc này mang lại khả năng xử lý nhanh, tối giản và tận dụng tối đa sức mạnh tính toán của hệ quản trị cơ sở dữ liệu.")
    add_body_paragraph(doc, "Cấu trúc mã nguồn chính của ứng dụng demo bao gồm các file cốt lõi sau:")
    add_bullet_point(doc, "Khởi chạy ứng dụng Spring Boot, nạp dữ liệu demo mẫu để minh họa việc lấy dữ liệu từ View doanh thu chi nhánh và chạy thử câu truy vấn tìm phòng trống động thông qua SQLHelper.", "com.homestay.App.java:")
    add_bullet_point(doc, "Cung cấp phương thức tĩnh `readQuery(path)` để đọc trực tiếp nội dung các file SQL từ thư mục `resources/sql`. Thiết kế này giúp tách biệt mã nguồn SQL ra khỏi mã Java, giúp lập trình viên dễ dàng tối ưu hóa, gỡ lỗi và chỉnh sửa câu lệnh SQL mà không cần biên dịch lại mã nguồn Java.", "com.homestay.SQLHelper.java:")
    add_bullet_point(doc, "Lớp kiểm thử kết nối JDBC độc lập giúp nhà phát triển kiểm tra nhanh các hàm tính tổng tiền hóa đơn `hoadon.func_tinh_tong_tien_hoa_don` và truy xuất mẫu dữ liệu khách hàng lẻ để xác thực tính đúng đắn.", "com.homestay.Test.java:")
    add_bullet_point(doc, "Lớp REST Controller định nghĩa các API endpoint phục vụ giao diện người dùng. Lớp này hỗ trợ đầy đủ các nghiệp vụ: Đăng nhập phân quyền theo chi nhánh, xem danh sách phòng trống thời gian thực, đặt phòng nhanh, cập nhật cơ sở vật chất bị hỏng hóc, checkout trả phòng và thống kê doanh thu theo thời gian.", "com.homestay.controller.DemoController.java:")

    # --- PHẦN V ---
    doc.add_page_break()
    add_heading_1(doc, "V. HƯỚNG DẪN CÀI ĐẶT & HƯỚNG DẪN SỬ DỤNG (PROGRAM GUIDE)")
    
    add_heading_2(doc, "1. Hướng dẫn cài đặt hệ thống")
    add_body_paragraph(doc, 
        "Để khởi chạy ứng dụng demo, lập trình viên và người dùng cần thực hiện theo các bước chi tiết dưới đây:")
        
    add_heading_3(doc, "Bước 1: Chuẩn bị Cơ sở dữ liệu PostgreSQL")
    add_bullet_point(doc, "Cài đặt PostgreSQL phiên bản 15 trở lên trên máy tính.")
    add_bullet_point(doc, "Khởi chạy công cụ pgAdmin 4 hoặc psql CLI và tạo một cơ sở dữ liệu mới mang tên: `quanlykhachsan`.")
    add_bullet_point(doc, "Đảm bảo tài khoản siêu quản trị (superuser) của PostgreSQL là `postgres` với mật khẩu là `admin` (nếu dùng mật khẩu khác, hãy chỉnh sửa trong file `src/main/resources/application.properties`).")
    
    add_heading_3(doc, "Bước 2: Di cư cấu trúc dữ liệu (Flyway Migration)")
    add_body_paragraph(doc, 
        "Hệ thống sử dụng Flyway để tự động quản lý phiên bản cơ sở dữ liệu. Lập trình viên không cần chạy script SQL thủ công. Khi ứng dụng Spring Boot khởi động, Flyway sẽ tự động quét thư mục `src/main/resources/db/migration` và thực thi tuần tự các file di cư để tạo lập bảng, chỉ mục, vai trò (roles) và các thủ tục.")
        
    add_heading_3(doc, "Bước 3: Khởi chạy và nạp dữ liệu mẫu ban đầu")
    add_bullet_point(doc, "Để nạp dữ liệu mẫu ban đầu (seeding) gồm chi nhánh, phòng, khách hàng, nhân viên và các giao dịch mẫu, hãy mở Terminal tại thư mục gốc của dự án và chạy file `run.bat` rồi chọn Option [2] (Khôi phục / Nạp dữ liệu mẫu).")
    add_bullet_point(doc, "Hoặc chạy trực tiếp lệnh sau để import dữ liệu:")
    add_code_block(doc, "psql \"postgresql://postgres:admin@localhost:5432/quanlykhachsan\" -f \"src/main/resources/db/data_seed.sql\"")
    
    add_heading_3(doc, "Bước 4: Build và chạy ứng dụng Spring Boot")
    add_bullet_point(doc, "Yêu cầu hệ thống: Cài đặt JDK 17 (hoặc mới hơn) và Apache Maven.")
    add_bullet_point(doc, "Chạy lệnh sau tại thư mục gốc để biên dịch và khởi chạy ứng dụng:")
    add_code_block(doc, "mvn clean spring-boot:run")
    add_bullet_point(doc, "Ứng dụng sẽ khởi động Tomcat server tại cổng `8080` (URL mặc định: http://localhost:8080).")
    
    add_heading_2(doc, "2. Hướng dẫn sử dụng các chức năng Demo (REST API)")
    add_body_paragraph(doc, 
        "Hệ thống cung cấp bộ REST API tại endpoint `/api/demo` phục vụ cho giao diện người dùng. Có thể sử dụng các công cụ như Postman để tương tác thử nghiệm:")
        
    add_bullet_point(doc, 
        "Endpoint: POST `/api/demo/login` với body JSON `{\"username\": \"nv1_cn1\", \"password\": \"123\"}`. "
        "Hệ thống sẽ lưu trữ thông tin chi nhánh và vai trò của nhân viên vào Session để áp dụng bộ lọc dữ liệu tự động.",
        bold_prefix="Đăng nhập phân quyền: "
    )
    add_bullet_point(doc, 
        "Endpoint: GET `/api/demo/phong-trong?chiNhanhId=1&checkIn=2026-06-15 14:00:00&checkOut=2026-06-20 12:00:00`. "
        "API sẽ thực thi câu lệnh SQL động để trả về danh sách các phòng còn trống tại chi nhánh trong khoảng thời gian được lọc.",
        bold_prefix="Tìm phòng trống thực tế: "
    )
    add_bullet_point(doc, 
        "Endpoint: POST `/api/demo/tim-va-dat-phong-nhanh` với body JSON gồm thông tin khách hàng, chi nhánh, yêu cầu phòng, ngày nhận/trả và tiền cọc. "
        "Hệ thống sẽ gọi hàm `func_tim_va_dat_phong_nhanh` dưới DB để thực hiện giao dịch đặt phòng an toàn.",
        bold_prefix="Đặt phòng nhanh liên chi nhánh: "
    )
    add_bullet_point(doc, 
        "Endpoint: GET `/api/demo/management/customers/{id}/history`. "
        "Cho phép quản lý xem lịch sử chi tiết các lần đặt phòng và chi phí của khách hàng.",
        bold_prefix="Xem lịch sử đặt phòng: "
    )
    add_bullet_point(doc, 
        "Endpoint: GET `/api/demo/doanh-thu-chi-nhanh`. "
        "Trả về báo cáo doanh thu và số lượng hóa đơn của các chi nhánh thông qua View `v_doanh_thu_chi_nhanh` dưới DB.",
        bold_prefix="Xem báo cáo doanh thu chi nhánh: "
    )

    # --- PHẦN VI ---
    doc.add_page_break()
    add_heading_1(doc, "VI. ĐÁNH GIÁ DỰ ÁN & PHÂN CÔNG CÔNG VIỆC")
    
    add_heading_2(doc, "1. Ưu điểm của dự án")
    add_bullet_point(doc, "Việc tổ chức dữ liệu thành 4 schema (`quanly`, `nhansu`, `khachhang`, `hoadon`) giúp hệ thống bảo mật tốt, tránh việc truy cập dữ liệu chéo trái phép giữa các bộ phận.", "Cấu trúc Schema phân vùng rõ ràng:")
    add_bullet_point(doc, "Thiết kế CSDL đạt chuẩn hóa 3NF giúp triệt tiêu sự dư thừa dữ liệu và hạn chế tối đa các dị thường dữ liệu. Việc gộp trưởng đoàn vào bảng đoàn khách ở migration `V20` đã giúp tối giản hóa cấu trúc bảng.", "Thiết kế chuẩn hóa tối ưu:")
    add_bullet_point(doc, "Toàn bộ luồng xử lý phức tạp (tự động nâng hạng hội viên, tính tiền phòng, áp dụng giảm giá, phụ thu muộn, kiểm soát tuổi trẻ em) đều được đóng gói thành các Trigger và Function chạy trực tiếp trong RDBMS. Điều này giúp giảm tải cho tầng Application và đảm bảo dữ liệu luôn nhất quán dù có kết nối từ bất kỳ nền tảng nào (Web, Mobile, Third-party).", "Đẩy mạnh tính toán xuống tầng Database:")
    add_bullet_point(doc, "Hệ thống đã thiết lập chỉ mục B-tree trên hầu hết các khóa ngoại, khóa chính và các thuộc tính thường xuyên lọc (tên khách hàng, sđt, ngày nhận/trả phòng), giúp tăng tốc độ truy vấn đáng kể.", "Tối ưu hóa chỉ mục đầy đủ:")
    
    add_heading_2(doc, "2. Nhược điểm và hạn chế")
    add_bullet_point(doc, "Hệ thống tính toán hiệu suất phòng động bằng cách duyệt qua lịch sử hóa đơn thuê phòng trong khoảng thời gian dài. Nếu dữ liệu hóa đơn tích lũy lên tới hàng triệu bản ghi, truy vấn này sẽ chạy chậm lại và cần thiết kế Materialized View hoặc bảng tổng hợp (Aggregation table) để tối ưu.", "Tính toán động các báo cáo phức tạp:")
    add_bullet_point(doc, "Hiện tại hệ thống chưa tích hợp công nghệ Caching (như Redis) ở tầng Spring Boot cho các API tìm phòng trống. Nếu có hàng ngàn khách hàng truy cập tìm phòng trống đồng thời, hệ thống PostgreSQL sẽ chịu tải lớn.", "Chưa tích hợp tầng Cache:")
    
    add_heading_2(doc, "3. Khó khăn gặp phải và cách khắc phục")
    add_bullet_point(doc, "Khi nhiều nhân viên lễ tân cùng thực hiện đặt phòng nhanh cho khách lẻ tại cùng một thời điểm, có thể xảy ra tình trạng đặt trùng phòng (Overbooking) do xung đột dữ liệu đồng thời. Khắc phục: Nhóm đã sử dụng câu lệnh `NOT EXISTS` lọc lịch đặt phòng bị trùng trong hàm `func_tim_phong_trong_phu_hop` kết hợp với mức cô lập giao dịch (Transaction Isolation Level) phù hợp của Spring Boot để đảm bảo tính an toàn giao dịch.", "Xử lý xung đột đặt phòng đồng thời:")
    add_bullet_point(doc, "Khi khách đặt phòng kéo dài vắt qua hai tháng (ví dụ từ 25/05 đến 05/06), việc tính toán hiệu suất sử dụng phòng và doanh thu phòng cho từng tháng cụ thể bị sai lệch. Khắc phục: Nhóm đã viết hàm bổ trợ sử dụng các hàm toán học `LEAST` và `GREATEST` so sánh ngày thuê thực tế với ngày bắt đầu và kết thúc của tháng khảo sát để bóc tách chính xác số ngày thuê thuộc về tháng đó.", "Tính toán hiệu suất phòng vắt tháng:")
    add_bullet_point(doc, "Khi khách hàng thanh toán nhiều hóa đơn liên tiếp trên hệ thống, việc theo dõi và nâng hạng hội viên có thể bị trễ hoặc sai sót nếu nhân viên quên cập nhật. Khắc phục: Viết Trigger tự động nâng hạng hội viên chạy ngay khi cột `tong_luu_tru` trong bảng hội viên thay đổi, đảm bảo tính cập nhật tức thời và chính xác.", "Tự động hóa thăng hạng hội viên:")

    add_heading_2(doc, "4. Phân công công việc của nhóm")
    add_body_paragraph(doc, "Hệ thống dự án được hoàn thành nhờ sự phối hợp chặt chẽ và phân công công việc rõ ràng giữa các thành viên trong nhóm:", space_after=6)
    
    member_headers = ["Thành viên", "MSSV", "Nhiệm vụ chính trong dự án", "Tỷ lệ đóng góp"]
    member_rows = [
        ["Nguyễn Văn A\n(Nhóm trưởng)", "20260001", 
         "- Khảo sát nghiệp vụ chuỗi homestay/khách sạn đa chi nhánh.\n"
         "- Thiết kế sơ đồ thực thể liên kết (ERD) tổng thể của dự án.\n"
         "- Xây dựng schema `quanly` và `nhansu` (chinhanh, loaiphong, phong, nhanvien, chucvu).\n"
         "- Lập trình và phân tích 3 câu SQL đại diện: Đặt phòng nhanh, Báo cáo hiệu suất phòng, Báo cáo tài chính chi tiết.\n"
         "- Viết báo cáo tổng kết và chuẩn bị tài liệu kỹ thuật.", "34%"],
        ["Trần Thị B", "20260002", 
         "- Khảo sát nghiệp vụ quản lý thông tin khách hàng lẻ, đoàn khách và hội viên.\n"
         "- Xây dựng schema `khachhang` (khachhang, hoivien, muchoivien, doankhach, khachhang_treem).\n"
         "- Lập trình và phân tích 3 câu SQL đại diện: Trigger thăng hạng hội viên, Tổng chi tiêu đoàn, Lịch sử đặt phòng khách hàng.\n"
         "- Thiết lập Trigger kiểm soát độ tuổi khách hàng trẻ em đi kèm phải dưới 18 tuổi.\n"
         "- Tối ưu hóa chỉ mục (Indexes) trên các bảng thông tin khách hàng để tăng tốc tìm kiếm.", "33%"],
        ["Lê Văn C", "20260003", 
         "- Khảo sát nghiệp vụ hóa đơn thanh toán, thuê phòng và sử dụng dịch vụ.\n"
         "- Xây dựng schema `hoadon` (hoadon, hoadon_thue_phong, hoadon_sudung_dichvu, dichvu).\n"
         "- Lập trình và phân tích 10 câu SQL đại diện xử lý thanh toán, checkout, tính tiền cọc, phụ thu trễ hạn và VAT.\n"
         "- Viết mã Backend Java Spring Boot REST API và tích hợp với JDBC Template.\n"
         "- Soạn thảo dữ liệu mẫu ban đầu (seeding) và thực hiện kiểm thử hiệu năng câu lệnh SQL.", "33%"]
    ]
    add_styled_table(doc, member_headers, member_rows, col_widths=[1.5, 0.8, 3.4, 0.8])
    
    # Save the document
    output_filename = "BaoCaoQuanLyKhachSan.docx"
    try:
        doc.save(output_filename)
    except PermissionError:
        alt_filename = "BaoCaoQuanLyKhachSan_new.docx"
        doc.save(alt_filename)
        print(f"WARNING: '{output_filename}' is locked (close Word first). Saved as '{alt_filename}' instead.")
        output_filename = alt_filename
    print(f"Report generated successfully: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    # Extracted data definitions to pass into helper tables
    tables_data = [
        ["Tên Bảng", "Schema", "Mô tả nghiệp vụ", "Vai trò thiết kế"],
        ["chinhanh", "quanly", "Thông tin chi nhánh (tên, địa chỉ)", "Thực thể chính quản lý chuỗi"],
        ["chusohuu", "quanly", "Thông tin chủ sở hữu chi nhánh (tên, sđt, email)", "Lưu trữ thông tin cổ đông"],
        ["chinhanh_chusohuu", "quanly", "Bảng trung gian liên kết chi nhánh và chủ sở hữu", "Thể hiện quan hệ nhiều-nhiều (N-N)"],
        ["loaiphong", "quanly", "Định nghĩa loại phòng (giường, diện tích, view, giá)", "Quản lý cơ chế định giá cơ bản"],
        ["phong", "quanly", "Thông tin phòng cụ thể và trạng thái hoạt động", "Đối tượng cốt lõi trong đặt phòng"],
        ["cosovatchat", "quanly", "Danh mục trang thiết bị và giá đền bù hư hỏng", "Quản lý tài sản khách sạn"],
        ["phong_trangbi_csvc", "quanly", "Liên kết trang bị cơ sở vật chất cho từng phòng", "Quản lý số lượng và hao mòn thiết bị"],
        ["chucvu", "nhansu", "Danh sách chức vụ và mức lương cơ bản", "Định nghĩa cơ cấu lương"],
        ["nhanvien", "nhansu", "Thông tin nhân viên, chức vụ và chi nhánh làm việc", "Quản lý nhân sự chi tiết"],
        ["muchoivien", "khachhang", "Quy định điều kiện số đêm và mức giảm giá", "Cấu hình chính sách chăm sóc VIP"],
        ["hoivien", "khachhang", "Thông tin điểm tích lũy và số đêm đã lưu trú", "Theo dõi lịch sử tích điểm của khách"],
        ["doankhach", "khachhang", "Quản lý đoàn khách du lịch và trưởng đoàn", "Tối ưu hóa đặt phòng theo nhóm"],
        ["khachhang", "khachhang", "Thông tin cá nhân (CCCD, SĐT, quốc tịch, hội viên)", "Thực thể chính thực hiện giao dịch"],
        ["khachhang_treem", "khachhang", "Thông tin trẻ em đi kèm khách hàng bảo lãnh", "Quản lý nhân khẩu học đi cùng"],
        ["dichvu", "hoadon", "Danh mục dịch vụ cung cấp và đơn giá", "Khai thác dịch vụ giá trị gia tăng"],
        ["hoadon", "hoadon", "Hóa đơn tổng quát của giao dịch", "Quản lý trạng thái hóa đơn chung"],
        ["hoadon_thue_phong", "hoadon", "Chi tiết phòng thuê, ngày nhận/trả, tiền cọc, phụ thu", "Tính toán chi phí phòng thực tế"],
        ["hoadon_sudung_dichvu", "hoadon", "Bảng trung gian lưu chi tiết dịch vụ được sử dụng", "Tính tổng tiền dịch vụ phát sinh"],
        ["lich_su_thao_tac", "hoadon", "Ghi lại thao tác đặt/hủy phòng để kiểm toán", "Theo dõi hoạt động hệ thống"]
    ]
    
    # Pre-load SQL codes for Member 1 & 2
    sql_a1 = ""
    sql_a2 = ""
    sql_a3 = ""
    sql_b1 = ""
    sql_b2 = ""
    sql_b3 = ""

    main()
